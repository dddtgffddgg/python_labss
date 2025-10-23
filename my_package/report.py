from docx import Document


def create_report(filename, shape_type, parameters, area):
    document = Document()
    
    document.add_heading("Shape Area Report", level=1)
    
    document.add_paragraph(f"Shape Type: {shape_type}")
    
    if shape_type == "triangle":
        base, height = parameters
        document.add_paragraph(f"Base: {base}")
        document.add_paragraph(f"Height: {height}")
    elif shape_type == "rectangle":
        length, width = parameters
        document.add_paragraph(f"Length: {length}")
        document.add_paragraph(f"Width: {width}")
    elif shape_type == "trapezoid":
        upper_base, lower_base, height = parameters
        document.add_paragraph(f"Upper Base: {upper_base}")
        document.add_paragraph(f"Lower Base: {lower_base}")
        document.add_paragraph(f"Height: {height}")
    
    document.add_paragraph(f"Area: {area}")
    
    document.save(filename)
