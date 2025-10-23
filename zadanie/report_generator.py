import math
from docx import Document

def generate_report(base, height, side_a, side_b):
    try:
        base = float(base)
        height = float(height)
        side_a = float(side_a)
        side_b = float(side_b)

        triangle_area = 0.5 * base * height
        rectangle_area = base * height
        trapezoid_area = 0.5 * (side_a + side_b) * height

        triangle_inscribed_radius = triangle_area / (base + height + math.sqrt(base**2 + height**2))
        triangle_circumscribed_radius = (base * height) / (math.sqrt(base**2 + height**2))
        rectangle_inscribed_radius = min(base, height) / 2
        rectangle_circumscribed_radius = math.sqrt(base**2 + height**2) / 2
        trapezoid_inscribed_radius = trapezoid_area / (side_a + side_b)
        trapezoid_circumscribed_radius = (side_a + side_b) / 4

        document = Document()
        document.add_heading('Площади фигур и радиусы описанных и вписанных окружностей', level=3)

        document.add_paragraph(f"Треугольник: {triangle_area}")
        document.add_paragraph(f"Прямоугольник: {rectangle_area}")
        document.add_paragraph(f"Трапеция: {trapezoid_area}")
        document.add_paragraph(f"Радиус вписанной окружности треугольника: {triangle_inscribed_radius}")
        document.add_paragraph(f"Радиус описанной окружности треугольника: {triangle_circumscribed_radius}")
        document.add_paragraph(f"Радиус вписанной окружности прямоугольника: {rectangle_inscribed_radius}")
        document.add_paragraph(f"Радиус описанной окружности прямоугольника: {rectangle_circumscribed_radius}")
        document.add_paragraph(f"Радиус вписанной окружности трапеции: {trapezoid_inscribed_radius}")
        document.add_paragraph(f"Радиус описанной окружности трапеции: {trapezoid_circumscribed_radius}")

        document.save('отчет.docx')

    except ValueError:
        print("Пожалуйста, введите числовые значения для параметров.")
