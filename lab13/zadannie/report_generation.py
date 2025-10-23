from docx import Document
from shapes import Rectangle, Triangle, Trapezoid


def generate_report(shapes):
    doc = Document()
    doc.add_heading('Shape Areas', level=1)

    for shape, area in shapes.items():
        if shape == 'Rectangle':
            doc.add_paragraph(f"Rectangle with base: {shape.base}, height: {shape.height}")
            doc.add_paragraph(f"Area: {area}")
        elif shape == 'Triangle':
            doc.add_paragraph(f"Triangle with base: {shape.base}, height: {shape.height}")
            doc.add_paragraph(f"Area: {area}")
            doc.add_paragraph(f"Circumscribed Circle Radius: {shape.calculate_circumscribed_circle_radius()}")
            doc.add_paragraph(f"Inscribed Circle Radius: {shape.calculate_inscribed_circle_radius()}")
        elif shape == 'Trapezoid':
            doc.add_paragraph(f"Trapezoid with base: {shape.base}, height: {shape.height}, side_a: {shape.side_a}, side_b: {shape.side_b}")
            doc.add_paragraph(f"Area: {area}")
            doc.add_paragraph(f"Circumscribed Circle Radius: {shape.calculate_circumscribed_circle_radius()}")
            doc.add_paragraph(f"Inscribed Circle Radius: {shape.calculate_inscribed_circle_radius()}")

    doc.save('otchet.docx')

