import appJar as aj
from docx import Document
from tkinter import filedialog

def calculate_triangle_area(base, height):
    return 0.5 * base * height

def calculate_triangle_perimeter(base, height):
    return base + (2 * (base**2 + height**2)**0.5)

def calculate_rectangle_area(length, width):
    return length * width

def calculate_rectangle_perimeter(length, width):
    return 2 * (length + width)

def calculate_trapezoid_area(top_base, bottom_base, height):
    return 0.5 * (top_base + bottom_base) * height

def calculate_trapezoid_perimeter(top_base, bottom_base, height):
    side = ((bottom_base - top_base) / 2)**2 + height**2
    return top_base + bottom_base + (2 * side)**0.5

def calculate_inscribed_circle_area(radius):
    return 3.14159 * radius**2

def calculate_inscribed_circle_perimeter(radius):
    return 2 * 3.14159 * radius

def calculate_circumscribed_circle_area(radius):
    return 4 * 3.14159 * radius**2

def calculate_circumscribed_circle_perimeter(radius):
    return 2 * 3.14159 * radius

def handle_button_press(button):
    if button == "Calculate Triangle":
        base = float(app.getEntry("Triangle Base"))
        height = float(app.getEntry("Triangle Height"))
        area = calculate_triangle_area(base, height)
        perimeter = calculate_triangle_perimeter(base, height)
        app.setLabel("Triangle Area", area)
        app.setLabel("Triangle Perimeter", perimeter)
        results.append("Triangle Area: " + str(area))
        results.append("Triangle Perimeter: " + str(perimeter))
    elif button == "Calculate Rectangle":
        length = float(app.getEntry("Rectangle Length"))
        width = float(app.getEntry("Rectangle Width"))
        area = calculate_rectangle_area(length, width)
        perimeter = calculate_rectangle_perimeter(length, width)
        app.setLabel("Rectangle Area", area)
        app.setLabel("Rectangle Perimeter", perimeter)
        results.append("Rectangle Area: " + str(area))
        results.append("Rectangle Perimeter: " + str(perimeter))
    elif button == "Calculate Trapezoid":
        top_base = float(app.getEntry("Trapezoid Top Base"))
        bottom_base = float(app.getEntry("Trapezoid Bottom Base"))
        height = float(app.getEntry("Trapezoid Height"))
        area = calculate_trapezoid_area(top_base, bottom_base, height)
        perimeter = calculate_trapezoid_perimeter(top_base, bottom_base, height)
        app.setLabel("Trapezoid Area", area)
        app.setLabel("Trapezoid Perimeter", perimeter)
        results.append("Trapezoid Area: " + str(area))
        results.append("Trapezoid Perimeter: " + str(perimeter))
    elif button == "Calculate Inscribed Circle":
        radius = float(app.getEntry("Inscribed Circle Radius"))
        area = calculate_inscribed_circle_area(radius)
        perimeter = calculate_inscribed_circle_perimeter(radius)
        app.setLabel("Inscribed Circle Area", area)
        app.setLabel("Inscribed Circle Perimeter", perimeter)
        results.append("Inscribed Circle Area: " + str(area))
        results.append("Inscribed Circle Perimeter: " + str(perimeter))
    elif button == "Calculate Circumscribed Circle":
        radius = float(app.getEntry("Circumscribed Circle Radius"))
        area = calculate_circumscribed_circle_area(radius)
        perimeter = calculate_circumscribed_circle_perimeter(radius)
        app.setLabel("Circumscribed Circle Area", area)
        app.setLabel("Circumscribed Circle Perimeter", perimeter)
        results.append("Circumscribed Circle Area: " + str(area))
        results.append("Circumscribed Circle Perimeter: " + str(perimeter))

def save_results_to_docx(results):
    if results:
        doc = Document()
        for result in results:
            doc.add_paragraph(result)
            doc.add_paragraph('')  # Add empty paragraph between results

        file_path = filedialog.asksaveasfilename(defaultextension=".docx")
        if file_path:
            doc.save(file_path)

app = aj.gui("Geometry Calculator")
app.setSticky("ew")

app.addLabelEntry("Triangle Base")
app.addLabelEntry("Triangle Height")
app.addButton("Calculate Triangle", handle_button_press)
app.addLabel("Triangle Area", "Area: ")
app.addLabel("Triangle Perimeter", "Perimeter: ")

app.addLabelEntry("Rectangle Length")
app.addLabelEntry("Rectangle Width")
app.addButton("Calculate Rectangle", handle_button_press)
app.addLabel("Rectangle Area", "Area: ")
app.addLabel("Rectangle Perimeter", "Perimeter: ")

app.addLabelEntry("Trapezoid Top Base")
app.addLabelEntry("Trapezoid Bottom Base")
app.addLabelEntry("Trapezoid Height")
app.addButton("Calculate Trapezoid", handle_button_press)
app.addLabel("Trapezoid Area", "Area: ")
app.addLabel("Trapezoid Perimeter", "Perimeter: ")

app.addLabelEntry("Inscribed Circle Radius")
app.addButton("Calculate Inscribed Circle", handle_button_press)
app.addLabel("Inscribed Circle Area", "Area: ")
app.addLabel("Inscribed Circle Perimeter", "Perimeter: ")

app.addLabelEntry("Circumscribed Circle Radius")
app.addButton("Calculate Circumscribed Circle", handle_button_press)
app.addLabel("Circumscribed Circle Area", "Area: ")
app.addLabel("Circumscribed Circle Perimeter", "Perimeter: ")

results = []  # List to store results

app.addButton("Save Results", lambda: save_results_to_docx(results))

app.go()
